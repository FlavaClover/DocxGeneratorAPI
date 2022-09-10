import io
from typing import List, Optional
from docx import Document
from docx.document import Document as DocumentObject
from src.exceptions import TemplateDoesNotContainsFields
from dataclasses import dataclass
import re
import datetime


@dataclass
class Field:
    name: str
    value: str


@dataclass
class Author:
    first_name: str
    second_name: str
    middle_name: Optional[str]


class User:
    def __init__(self, login: str, password: str, first_name: str, second_name: str, middle_name: Optional[str]):
        self.login = login
        self.password = password
        self.first_name = first_name
        self.second_name = second_name
        self.middle_name = middle_name
        self._templates = list()
        self._documents = list()

    @property
    def author(self) -> Author:
        return Author(first_name=self.first_name, second_name=self.second_name, middle_name=self.middle_name)

    def __repr__(self):
        return 'User({0}, {1}, {2})'.format(self.first_name, self.second_name, self.middle_name)

    def __eq__(self, other):
        if not isinstance(other, User):
            return False

        return self.login == other.login and \
            self.password == other.password and \
            self.first_name == other.first_name and \
            self.second_name == other.second_name and \
            self.middle_name == self.middle_name


class Docx:
    def __init__(self, name: str, content: bytes | str, author: User,
                 created_datetime: datetime.datetime = datetime.datetime.now()):
        if isinstance(content, str):
            with open(content, 'rb') as file:
                content = file.read()

        self.content = content
        self._bytes_io = io.BytesIO(content)
        self.document: DocumentObject = Document(self._bytes_io)
        self._name = name
        self._created_datetime = created_datetime
        self._author = author
        self.user_id = None if not hasattr(author, 'id') else author.id

    @property
    def created_datetime(self):
        return self._created_datetime

    @property
    def name(self):
        return self._name + '.docx'

    def __hash__(self):
        return hash(self.content)

    def __repr__(self):
        return self.name


class Template(Docx):
    def __init__(self, name: str, content: bytes | str, author: User,
                 created_datetime: datetime.datetime = datetime.datetime.now(), ):
        super().__init__(name, content, author, created_datetime=created_datetime)

        if not self._is_have_template_fields():
            raise TemplateDoesNotContainsFields('Template must be contains any fields')

    def _is_have_template_fields(self) -> bool:
        is_have = False

        paragraphs = [p.text for p in self.paragraphs if p.text != '']

        for paragraph in paragraphs:
            print(paragraph)
            if len(re.findall(r'___.+___', paragraph)) > 0:
                is_have = True
                break
        return is_have

    def _get_cells_paragraphs(self) -> List:
        paragraphs = []

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs += cell.paragraphs

        return paragraphs

    def _get_text_paragraphs(self) -> List:
        paragraphs = []

        for paragraph in self.document.paragraphs:
            paragraphs.append(paragraph)

        return paragraphs

    @property
    def fields(self) -> List:
        fields = []
        for paragraph in self.paragraphs:
            fields.extend(re.findall('___.+___', paragraph.text))

        return fields

    @property
    def paragraphs(self) -> List:
        return self._get_text_paragraphs() + self._get_cells_paragraphs()

    def __hash__(self):
        return super().__hash__()

    def __repr__(self):
        return self.name
