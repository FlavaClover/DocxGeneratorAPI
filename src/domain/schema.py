from pydantic import BaseModel
from pydantic.dataclasses import dataclass
from typing import Optional, List
import io
import datetime
from docx import Document
from docx.document import Document as DocumentObject
from src.domain.exceptions import TemplateDoesNotContainsFields, TemplateDoesNotContainsSpecificField
import re


@dataclass
class Author:
    first_name: str
    second_name: str
    middle_name: Optional[str]


@dataclass
class User:
    login: str
    password: str
    first_name: str
    second_name: str
    middle_name: str | None

    def __init__(self, login: str, password: str, first_name: str, second_name: str, middle_name: Optional[str]):
        self.login = login
        self.password = password
        self.first_name = first_name
        self.second_name = second_name
        self.middle_name = middle_name

    @property
    def author(self) -> Author:
        return Author(first_name=self.first_name, second_name=self.second_name, middle_name=self.middle_name)

    def __repr__(self):
        return 'User({0}, {1}, {2})'.format(self.first_name, self.second_name, self.middle_name)

    def __eq__(self, other):
        if not isinstance(other, User):
            return False

        return self.login == other.login and \
            self.password == other.password and \
            self.first_name == other.first_name and \
            self.second_name == other.second_name and \
            self.middle_name == self.middle_name

    def add_template(self, *templates):
        self._templates.extend(templates)

    def remove_template(self, template):
        self._templates.remove(template)

    def add_document(self, document):
        self._documents.append(document)


@dataclass
class Docx:
    def __init__(self, name: str, content: bytes | str,
                 created_datetime: datetime.datetime = datetime.datetime.now()):
        if isinstance(content, str):
            with open(content, 'rb') as file:
                content = file.read()

        self.content = content
        self._bytes_io = io.BytesIO(content)
        self.document: DocumentObject = Document(self._bytes_io)
        self._name = name
        self._created_datetime = created_datetime

    @property
    def created_datetime(self):
        return self._created_datetime

    @property
    def name(self):
        return self._name + '.docx'

    def __hash__(self):
        return hash(self.content)

    def __repr__(self):
        return self.name


@dataclass
class Template(Docx):
    def __init__(self, name: str, content: bytes | str,
                 created_datetime: datetime.datetime = datetime.datetime.now(), auto_fill_fields=True):
        super().__init__(name, content, created_datetime=created_datetime)
        self.fields = self._fields() if auto_fill_fields else []
        if not self._is_have_template_fields():
            raise TemplateDoesNotContainsFields('Template must be contains any fields')

    def _is_have_template_fields(self) -> bool:
        is_have = False

        paragraphs = [p.text for p in self.paragraphs if p.text != '']

        for paragraph in paragraphs:
            if len(re.findall(r'___.+___', paragraph)) > 0:
                is_have = True
                break
        return is_have

    def _get_cells_paragraphs(self) -> List:
        paragraphs = []

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs += cell.paragraphs

        return paragraphs

    def _get_text_paragraphs(self) -> List:
        paragraphs = []

        for paragraph in self.document.paragraphs:
            paragraphs.append(paragraph)

        return paragraphs

    def _fields(self) -> List:
        fields = []
        for paragraph in self.paragraphs:
            fields.extend(re.findall('___.+___', paragraph.text))

        return list(map(lambda x: TemplateField(self, x), fields))

    @property
    def template_fields(self):
        return self.fields

    def add_field(self, field_name):
        field = TemplateField(self, field_name)
        if field not in self._fields():
            raise TemplateDoesNotContainsSpecificField('Template does not contains this field')
        else:
            self.fields.append(field)

    def remove_field(self, field_name):
        self.fields.remove(TemplateField(self, field_name))

    def auto_fill_field(self):
        self.fields = self._fields()

    @property
    def paragraphs(self) -> List:
        return self._get_text_paragraphs() + self._get_cells_paragraphs()

    def __hash__(self):
        return super().__hash__()

    def __repr__(self):
        return self.name


@dataclass
class TemplateField:
    def __init__(self, template: Template, name: str):
        self.template = template
        self.template_id = None if not hasattr(template, 'id') else template.id
        self.name = name

    def __eq__(self, other):
        return self.name == other.name and self.template == other.template

    def __repr__(self):
        return '{0}: {1}'.format(self.template_id, self.name)